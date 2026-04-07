import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()
    
    # helper for adding a block of text
    def add_text(text, bold=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Titel des Job Postings")
    run.bold = True
    run.font.size = Pt(14)

    # Table for EKP and SVS
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = 'angefragt'
    hdr_cells[2].text = 'falls abweichend'
    
    # EKP Row
    row = table.rows[1].cells
    row[0].text = 'Einkaufskurzprofil (EKP)'
    row[1].text = 'X|YYY|XXX|Z'
    row[2].text = 'X|YYY|XXX|Z'
    
    # SVS Row
    row = table.rows[2].cells
    row[0].text = 'Stundenverrechnungssatz (SVS)'
    row[1].text = '€'
    row[2].text = '€'

    doc.add_paragraph() # Spacer

    add_text("Möglicher Starttermin (verfügbar ab / Kündigungsfrist): 15.04.2026")

    doc.add_paragraph() # Spacer

    # Education loop
    add_text("{% if education.higher_education %}{%tr for he in education.higher_education %}{{ he.years }}{{ he.institution }} – {{ he.field }}{%tr endfor %}{% endif %}", bold=True)

    doc.add_paragraph() # Spacer
    add_text("Beruferfahrung", bold=True, size=12)
    doc.add_paragraph() # Spacer

    # Employment logic
    emp_text = """{%tr for job in employment_history %}	Arbeitgeber: {{ job.employer }}Position: {{ job.position }}
Tätigkeit: {% for d in job.duties %}{{ d }}{% endfor %}
{{ job.start_date }} – {{ job.end_date }}
		Arbeitgeber: {{ job.employer }}Position: {{ job.position }}
Tätigkeit: {% for d in job.duties %}{{ d }}{% endfor %}
{%tr endfor %}"""
    add_text(emp_text)

    doc.add_paragraph() # Spacer

    # Higher Education Block
    add_text("{% if education.higher_education %}", size=10)
    add_text("Bildungseinrichtung:", bold=True)
    add_text("{%tr for he in education.higher_education %}")
    add_text("{{ he.years }}: {{ he.institution }} – {{ he.field }}")
    add_text("{%tr endfor %}")
    add_text("{% endif %}", size=10)

    doc.add_paragraph() # Spacer

    # Further Training Block
    add_text("{% if education.further_training %}", size=10)
    add_text("WEITERBILDUNG:", bold=True)
    add_text("{%tr for ft in education.further_training %}")
    add_text("{{ ft.years }}: {{ ft.institution }} – {{ ft.field }}")
    add_text("{%tr endfor %}")
    add_text("{% endif %}", size=10)

    doc.add_paragraph() # Spacer

    # Sudor (Welder) Skills
    add_text("{% if job_type == \"Sudor\" %}", size=10)
    add_text("Fähigkeiten", bold=True)
    skills = [
        "Schweißer (MIG/MAG 135)",
        "Handwerkliches Geschick: Sicherer Umgang mit Schweißgeräten, Brennern und Handwerkzeugen",
        "Installation: Montage und Heften von Schweißbaugruppen nach Zeichnung und Schweißplan",
        "Demontage: Trennen von Metallverbindungen mittels Brennschneider, Trennschleifer und Bohrmaschine",
        "Blechbearbeitung: Zuschneiden, Anfasen, Richten und Vorbereiten von Blechen und Profilen",
        "Montieren von Stahlkonstruktionen: Zusammensetzen, Ausrichten und Fixieren von Schweißteilen und Träger",
        "MIG-MAG 135, 136, 138: ",
        "Arbeit mit Robotern: Bedienung, Programmierung und Überwachung von Schweißrobotern (MIG/MAG-Automation)"
    ]
    for s in skills:
        doc.add_paragraph(s, style='List Bullet')
    add_text("{% endif %}", size=10)

    doc.add_paragraph() # Spacer

    # Language Skills
    add_text("Sprachkenntnisse:", bold=True)
    add_text("{%tr for l in language_skills %}")
    add_text("{{ l.language }}: {{ l.level }}")
    add_text("{%tr endfor %}")

    # Ensure resources directory exists
    output_dir = "app/resources/templates"
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, "Standard_Vorlage.docx")
    doc.save(output_path)
    print(f"Template created at: {output_path}")

if __name__ == "__main__":
    create_template()
