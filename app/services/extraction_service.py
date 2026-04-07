import json
import logging
import time
import re
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from app.models.candidate import Candidate
from app.models.employment import EmploymentRecord
from app.models.skill import SkillRecord
from app.models.language import LanguageRecord
from app.domain.normalization import strip_gdpr_data, detect_gaps
from app.ai.provider_manager import ProviderManager

logger = logging.getLogger(__name__)


class ExtractionService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.provider_manager = ProviderManager()

    def _extract_json(self, text: str) -> str:
        """Robustly extract JSON block from AI string response.
        
        Handles: raw JSON, markdown code fences, leading/trailing text.
        """
        # 1. Try to strip markdown code fences first
        fence_pattern = r"```(?:json)?\s*\n?([\s\S]*?)\n?\s*```"
        fence_match = re.search(fence_pattern, text)
        if fence_match:
            return fence_match.group(1).strip()
        
        # 2. Find JSON object — match balanced braces from first '{' to matching '}'
        first_brace = text.find('{')
        if first_brace == -1:
            raise ValueError("No JSON object found in AI response")
        
        depth = 0
        in_string = False
        escape = False
        end = -1
        for i in range(first_brace, len(text)):
            ch = text[i]
            if escape:
                escape = False
                continue
            if ch == '\\' and in_string:
                escape = True
                continue
            if ch == '"' and not escape:
                in_string = not in_string
                continue
            if in_string:
                continue
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    end = i
                    break
        
        if end == -1:
            raise ValueError("Unmatched braces in AI response")
        
        return text[first_brace:end + 1]

    async def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from PDF or DOCX file."""
        from pathlib import Path
        import aiofiles
        full_path = Path(file_path)
        if not full_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = full_path.suffix.lower()
        if suffix == ".pdf":
            import pdfplumber
            text_parts = []
            with pdfplumber.open(str(full_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts)
        elif suffix in (".docx", ".doc"):
            from docx import Document
            doc = Document(str(full_path))
            text_parts = []
            
            # Extract from main body paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())
            
            # Extract from all tables (critical for CVs formatted with tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                text_parts.append(p.text.strip())
                                
            return "\n".join(text_parts)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    async def extract_and_process(self, candidate_id: str):
        # Background tasks should use their own clean session context
        # If self.db was a request-scoped session, it might be closed.
        
        candidate = await self.db.get(Candidate, candidate_id)
        if not candidate:
            raise ValueError(f"Candidate {candidate_id} not found")

        candidate.status = "EXTRACTING"
        # Immediate save for status
        self.db.add(candidate)
        await self.db.commit()
        # Ensure we keep the object tracked
        await self.db.refresh(candidate)

        try:
            from app.config import settings
            from app.storage.local import LocalStorage
            storage = LocalStorage(settings.STORAGE_ROOT)
            
            # [LOG] File Metadata
            logger.info(f"DB_LOG: Metadata for {candidate_id} | File: {candidate.original_filename} | Path: {candidate.original_file_path}")
            
            raw_bytes = await storage.read(candidate.original_file_path)
            file_size = len(raw_bytes)
            logger.info(f"DB_LOG: Blob size: {file_size} bytes")

            import tempfile
            import os
            from pathlib import Path
            suffix = Path(candidate.original_filename).suffix
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(raw_bytes)
                tmp_path = tmp.name

            try:
                raw_text = await self.extract_text_from_file(tmp_path)
                text_len = len(raw_text)
                
                # [LOG POINT 1] Raw Extracted Text
                logger.info(f"DB_LOG: Raw Text Length: {text_len} chars")
                logger.info(f"DB_LOG: Raw Text Start: {raw_text[:800]}...")
                
            except Exception as e:
                logger.error(f"Failed to extract text: {e}")
                raise
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

            cleaned_text, removed_pii = strip_gdpr_data(raw_text)

            messages = [
                {"role": "system", "content": """You are a professional German recruitment CV parser. You extract structured data accurately from CV text.
CRITICAL RULES:
- Return ONLY valid JSON. No markdown, no explanations, no code fences.
- If a field cannot be found in the CV text, use null — NOT an empty string, NOT "Unknown", NOT "N/A".
- full_name MUST be the person's actual full name (e.g. "Max Mustermann"). If you cannot find a person name, set full_name to null.
- nationality MUST be the person's citizenship/country (e.g. "Deutsch", "Polnisch"). If not found, set to null.
- If the text appears to be a template with placeholder variables (like {{ name }}), set full_name and nationality to null."""},
                {"role": "user", "content": f"""Extract the following CV data into a structured JSON format. 
Fields: full_name (string or null), date_of_birth (YYYY-MM-DD or null), nationality (string or null), employment_history (array of objects), skills (array), languages (array of objects).
Employment history fields: company_name, job_title, start_date (YYYY-MM-DD), end_date, is_current (bool), description (list of strings).

CV TEXT:
{cleaned_text[:5000]}"""}
            ]

            start_time = time.time()
            try:
                raw_ai_res, provider_name, model = await self.provider_manager.try_extraction(messages, self.db)
                duration_ms = int((time.time() - start_time) * 1000)
                
                # [LOG POINT 2] Raw AI Response
                logger.info(f"DB_LOG: Provider: {provider_name} ({model})")
                logger.info(f"DB_LOG: Raw AI Response: {raw_ai_res}")
                
            except Exception as e:
                logger.error(f"AI Provider call failed: {e}")
                raise

            try:
                json_string = self._extract_json(raw_ai_res)
                result = json.loads(json_string)
                
                # [LOG POINT 3] Final Parsed JSON
                logger.info(f"DB_LOG: Final Parsed JSON: {json.dumps(result, indent=2, ensure_ascii=False)}")
                
                # Validate that result is a dict (not a list or string)
                if not isinstance(result, dict):
                    raise ValueError(f"AI returned {type(result).__name__} instead of JSON object: {str(result)[:200]}")
                
            except Exception as e:
                logger.error(f"Failed to parse JSON: {e} | Raw string output was: {raw_ai_res[:500]}")
                raise ValueError(f"AI returned invalid JSON: {e}")

            # Re-verify and update
            # Refresh again to be 100% sure we have an active instance
            await self.db.refresh(candidate)
            
            candidate.extraction_provider = provider_name
            candidate.extraction_model = model
            candidate.extraction_duration_ms = duration_ms

            # --- Extract identity fields with fallbacks ---
            # Try multiple key names the AI might use for the person's name
            full_name = (
                result.get("full_name")
                or result.get("name")
                or result.get("candidate_name")
                or result.get("person_name")
                or ""
            )
            # If the value is a dict (e.g., {"first": "Max", "last": "Mustermann"}), flatten it
            if isinstance(full_name, dict):
                parts = [full_name.get(k) for k in ("first_name", "last_name", "first", "last", "given_name", "family_name") if full_name.get(k)]
                full_name = " ".join(parts).strip()
            # If it's a list, join
            if isinstance(full_name, list):
                full_name = " ".join(str(x) for x in full_name if x).strip()
            
            # Clean up: strip whitespace, remove "Unknown"/"N/A" placeholders
            if isinstance(full_name, str):
                full_name = full_name.strip()
                if full_name.lower() in ("unknown", "n/a", "null", "none", "", "—", "-"):
                    full_name = ""
            
            candidate.full_name = full_name or None

            # Nationality with fallbacks
            nationality = (
                result.get("nationality")
                or result.get("citizenship")
                or result.get("country_of_origin")
                or ""
            )
            if isinstance(nationality, str):
                nationality = nationality.strip()
                if nationality.lower() in ("unknown", "n/a", "null", "none", "", "—", "-"):
                    nationality = ""
            candidate.nationality = nationality or None

            dob_str = result.get("date_of_birth")
            if dob_str:
                try:
                    candidate.date_of_birth = datetime.strptime(str(dob_str)[:10], "%Y-%m-%d").date()
                except Exception:
                    logger.warning(f"Failed to parse DOB: {dob_str}")

            # Persist Name changes immediately
            self.db.add(candidate)
            
            employment_data = result.get("employment_history", [])
            gaps = detect_gaps(employment_data)

            for emp in employment_data + gaps:
                def parse_date(d):
                    if not d or str(d).lower() in ("null", "none", "unknown"):
                        return None
                    try:
                        return datetime.strptime(str(d)[:10], "%Y-%m-%d").date()
                    except Exception:
                        return None

                record = EmploymentRecord(
                    candidate_id=candidate_id,
                    company_name=emp.get("company_name", "Unknown"),
                    job_title=emp.get("job_title", "Unknown"),
                    job_title_original=emp.get("job_title"),
                    start_date=parse_date(emp.get("start_date")),
                    end_date=parse_date(emp.get("end_date")),
                    is_current=emp.get("is_current", False),
                    description=emp.get("description", []),
                    is_gap_record=emp.get("is_gap_record", False),
                    gap_type=emp.get("gap_type"),
                    gap_note=emp.get("gap_note"),
                    inferred=emp.get("is_gap_record", False) or emp.get("inferred", False),
                    needs_review=True
                )
                self.db.add(record)

            for skill in result.get("skills", []):
                # Handle both dict-based and string-based skills from AI
                if isinstance(skill, str):
                    name = skill
                    cat = "SONSTIGES"
                    lvl = "KENNTNISSE"
                    ev = None
                else:
                    name = skill.get("name", skill.get("skill_name", ""))
                    cat = skill.get("category", "SONSTIGES")
                    lvl = skill.get("level", "KENNTNISSE")
                    ev = skill.get("evidence")

                record = SkillRecord(
                    candidate_id=candidate_id,
                    skill_name=name,
                    category=cat,
                    level=lvl,
                    evidence=ev,
                    inferred=True,
                    needs_review=True
                )
                self.db.add(record)

            for lang in result.get("languages", []):
                if isinstance(lang, str):
                    name = lang
                    raw = None
                    norm = "A1"
                else:
                    name = lang.get("language", "")
                    raw = lang.get("level_raw")
                    norm = lang.get("level_normalized", "A1")

                record = LanguageRecord(
                    candidate_id=candidate_id,
                    language=name,
                    level_raw=raw,
                    level_normalized=norm,
                    inferred=True,
                    needs_review=True
                )
                self.db.add(record)

            candidate.status = "EXTRACTED"
            candidate.needs_review = True
            
            # If no identity data was extracted at all, flag for review more prominently
            if not candidate.full_name and not nationality:
                candidate.needs_review = True
                logger.warning(
                    f"Candidate {candidate_id}: EXTRACTED but no identity data (name={candidate.full_name}, nationality={candidate.nationality}). "
                    f"CV may be a template or have no personal data."
                )
            
            # THE FINAL SAVE
            self.db.add(candidate)
            await self.db.commit()
            await self.db.refresh(candidate)
            logger.info(f"Candidate {candidate_id} successfully processed and saved.")

        except Exception as e:
            candidate.status = "FAILED"
            candidate.error_log = {"error": str(e), "timestamp": datetime.now().isoformat()}
            await self.db.commit()
            logger.error(f"Extraction flow failed for candidate {candidate_id}: {e}")
            raise
