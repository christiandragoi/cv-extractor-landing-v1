"""
Template Population Service for CV Extractor.

Reads a Jinja2-style Word template (.docx) and fills it with structured
candidate data extracted from a CV by AI.

The template uses placeholders like:
  {{ job.employer }}, {{ job.start_date }}, {{ l.language }}, etc.
  {%tr for job in employment_history %} ... {%tr endfor %}
"""

import json
import re
import copy
import logging
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


class TemplatePopulationService:
    """Fills a Word template with candidate data."""

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.template_doc = Document(template_path)
        self._font_name = "Calibri"
        self._font_size = Pt(11)

    def populate(self, data: dict) -> Document:
        """
        Populate the template with extracted candidate data.

        Args:
            data: Structured dict matching the expected schema:
                  - job_posting: {title, ekp, svs, start_date}
                  - employment_history: [{start_date, end_date, employer, position, duties}]
                  - education: {higher_education: [...], further_training: [...]}
                  - skills: {items: [{name, level}]}
                  - language_skills: [{language, level}]
                  - technical_skills: [{name, level}]

        Returns:
            A new Document object with all placeholders filled.
        """
        doc = Document(self.template_path)
        self._doc = doc

        # --- Table 0: Job Posting Header ---
        self._fill_job_posting_header(data.get("job_posting", {}))

        # --- Table 1: Berufserfahrung (Employment History) ---
        self._fill_employment_history(data.get("employment_history", []))

        # --- Table 2: Bildungseinrichtung (Education) ---
        higher_ed = data.get("education", {}).get("higher_education", [])
        self._fill_education_table(doc.tables[2], higher_ed)

        # --- Table 3: Weiterbildung (Further Training) ---
        further = data.get("education", {}).get("further_training", [])
        self._fill_education_table(doc.tables[3], further)

        # --- Table 4: Fähigkeiten / Technical Skills ---
        self._fill_skills_table(data.get("technical_skills", []))

        # --- Table 5: Sprachkenntnisse (Languages) ---
        self._fill_languages_table(data.get("language_skills", []))

        # --- Paragraphs: Remove Jinja2 conditionals ---
        self._clean_paragraphs(doc)

        # --- Clean leftover Jinja2 rows from ALL tables ---
        self._clean_template_rows(doc)

        # --- Clean static skills row if it exists (old template artifact) ---
        self._clean_static_skills_row(doc)

        return doc

    def save(self, data: dict, output_path: str):
        """Populate and save to file."""
        doc = self.populate(data)
        doc.save(output_path)
        logger.info(f"Populated template saved to: {output_path}")

    def get_bytes(self, data: dict) -> bytes:
        """Populate and return as bytes."""
        import io
        doc = self.populate(data)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # Internal methods
    # ------------------------------------------------------------------

    def _fill_job_posting_header(self, job_posting: dict):
        """Fill Table 0: EKP, SVS, Starttermin."""
        if not job_posting or len(self._doc.tables) < 1:
            return
        table = self._doc.tables[0]

        # Row 2, Col 1: EKP
        if job_posting.get("ekp"):
            self._set_cell_text(table.rows[2].cells[1], job_posting["ekp"])
        # Row 2, Col 2: EKP fallback
        if job_posting.get("ekp"):
            self._set_cell_text(table.rows[2].cells[2], job_posting["ekp"])

        # Row 3, Col 1-2: SVS
        if job_posting.get("svs"):
            self._set_cell_text(table.rows[3].cells[1], f"€ {job_posting['svs']}")
            self._set_cell_text(table.rows[3].cells[2], f"€ {job_posting['svs']}")

        # Row 4, Col 1-2: Starttermin
        if job_posting.get("start_date"):
            self._set_cell_text(table.rows[4].cells[1], job_posting["start_date"])
            self._set_cell_text(table.rows[4].cells[2], job_posting["start_date"])

        # Row 0: Title
        if job_posting.get("title"):
            # Clear and set the title cell
            cell = table.rows[0].cells[0]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            cell.paragraphs[0].runs[0].text = job_posting["title"] if cell.paragraphs[0].runs else None
            if not cell.paragraphs[0].runs:
                run = cell.paragraphs[0].add_run(job_posting["title"])
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = self._font_name

    def _fill_employment_history(self, jobs: list):
        """Fill Table 1: Berufserfahrung with dynamic rows."""
        if not jobs or len(self._doc.tables) < 2:
            return

        table = self._doc.tables[1]
        # Table 1, Row 0 is the header "Beruferfahrung"
        # Table 1, Rows 1-2 are template rows with Jinja2 syntax

        # Remove template rows (rows 1 and 2, which have Jinja2 placeholders)
        # We'll add fresh rows for each job

        # Collect the template row XML for cloning (row 2 has the format we want)
        template_row_xml = None
        if len(table.rows) > 2:
            template_row_xml = copy.deepcopy(table.rows[2]._tr)

        # Remove existing template rows (keep header row 0)
        rows_to_remove = []
        for i in range(1, len(table.rows)):
            rows_to_remove.append(table.rows[i]._tr)
        for tr in rows_to_remove:
            table._tbl.remove(tr)

        # Add a row for each job
        for idx, job in enumerate(jobs):
            if template_row_xml:
                new_tr = copy.deepcopy(template_row_xml)
                table._tbl.append(new_tr)

                # Now populate the cells
                last_row_idx = len(table.rows) - 1
                row = table.rows[last_row_idx]

                # Left cell: date range
                date_range = f"{job.get('start_date', '')} – {job.get('end_date', '')}"
                self._set_cell_text(row.cells[0], date_range)

                # Right cell: employer, position, duties
                parts = []
                if job.get("employer"):
                    parts.append(f"Arbeitgeber: {job['employer']}")
                if job.get("position"):
                    parts.append(f"Position: {job['position']}")
                if job.get("duties"):
                    duties_text = ", ".join(job["duties"]) if isinstance(job["duties"], list) else str(job["duties"])
                    parts.append(f"Tätigkeit: {duties_text}")
                self._set_cell_text(row.cells[1], "\n".join(parts))
            else:
                # Fallback: add new rows manually
                row = table.add_row()
                date_range = f"{job.get('start_date', '')} – {job.get('end_date', '')}"
                self._set_cell_text(row.cells[0], date_range)

                parts = []
                if job.get("employer"):
                    parts.append(f"Arbeitgeber: {job['employer']}")
                if job.get("position"):
                    parts.append(f"Position: {job['position']}")
                if job.get("duties"):
                    duties_text = ", ".join(job["duties"]) if isinstance(job["duties"], list) else str(job["duties"])
                    parts.append(f"Tätigkeit: {duties_text}")
                self._set_cell_text(row.cells[1], "\n".join(parts))

    def _fill_education_table(self, table, entries: list):
        """Fill Table 2 or 3: Education/Training."""
        if not entries:
            # Remove all rows except the header
            return

        # Remove template rows (keep row structure)
        rows_to_remove = []
        for i in range(0, len(table.rows)):
            text = table.rows[i].cells[0].text if table.rows[i].cells else ""
            if "{%" in text or "{{" in text:
                rows_to_remove.append(table.rows[i]._tr)
        for tr in rows_to_remove:
            table._tbl.remove(tr)

        # Get template row for cloning
        template_row = None
        if len(table.rows) > 1:
            template_row = copy.deepcopy(table.rows[1]._tr) if len(table.rows) > 1 else None

        for entry in entries:
            if template_row:
                new_tr = copy.deepcopy(template_row)
                table._tbl.append(new_tr)
                last_row = table.rows[len(table.rows) - 1]
                self._set_cell_text(last_row.cells[0], entry.get("years", ""))
                detail = f"{entry.get('institution', '')} – {entry.get('field', '')}"
                self._set_cell_text(last_row.cells[1], detail)
            else:
                row = table.add_row()
                self._set_cell_text(row.cells[0], entry.get("years", ""))
                detail = f"{entry.get('institution', '')} – {entry.get('field', '')}"
                self._set_cell_text(row.cells[1], detail)

    def _fill_skills_table(self, skills: list):
        """Fill Table 4: Technical skills with levels."""
        if not skills or len(self._doc.tables) < 5:
            return

        table = self._doc.tables[4]
        cell = table.rows[0].cells[0]

        # Clear ALL content in the entire table first (remove all rows, keep one)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # Clear the remaining row's cell completely
        for p in cell.paragraphs:
            p.clear()

        # Build skills list with bullet formatting
        skill_lines = []
        for skill in skills:
            name = skill.get("name", "")
            level = skill.get("level", "")
            if name and level:
                skill_lines.append(f"  {name}: {level}")
            elif name:
                skill_lines.append(f"  {name}")

        # Write to first paragraph of the cell
        if cell.paragraphs:
            first_p = cell.paragraphs[0]
            run = first_p.add_run("Sonstige Techniken / Schweißfähigkeiten:\n")
            run.bold = True
            run.font.name = self._font_name
            run.font.size = self._font_size

            for line in skill_lines:
                run = first_p.add_run(f"  - {line.strip()}\n")
                run.font.name = self._font_name
                run.font.size = self._font_size

    def _fill_languages_table(self, languages: list):
        """Fill Table 5: Language skills."""
        if not languages or len(self._doc.tables) < 6:
            return

        table = self._doc.tables[5]

        # Remove template rows
        rows_to_remove = []
        for i in range(len(table.rows)):
            text = table.rows[i].cells[0].text if table.rows[i].cells else ""
            if "{%" in text or "{{" in text:
                rows_to_remove.append(table.rows[i]._tr)
        for tr in rows_to_remove:
            table._tbl.remove(tr)

        # Get template row for cloning (second row has the format)
        template_row = None
        for row in table.rows:
            if "{{" in (row.cells[0].text if row.cells else ""):
                template_row = copy.deepcopy(row._tr)
                break

        for lang in languages:
            if template_row:
                new_tr = copy.deepcopy(template_row)
                table._tbl.append(new_tr)
                last_row = table.rows[len(table.rows) - 1]
                self._set_cell_text(last_row.cells[0], lang.get("language", ""))
                self._set_cell_text(last_row.cells[1], lang.get("level", ""))
            else:
                row = table.add_row()
                self._set_cell_text(row.cells[0], lang.get("language", ""))
                self._set_cell_text(row.cells[1], lang.get("level", ""))

    def _clean_template_rows(self, doc: Document):
        """Remove any table rows that still contain Jinja2 placeholders."""
        jinja_pattern = re.compile(r'\{%.*?%\}|\{\{.*?\}\}')
        for table in doc.tables:
            rows_to_remove = []
            for row in table.rows:
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text
                if jinja_pattern.search(row_text):
                    rows_to_remove.append(row._tr)
            for tr in rows_to_remove:
                table._tbl.remove(tr)

    def _clean_static_skills_row(self, doc: Document):
        """Remove the old static skills row that has no Jinja2 markers."""
        for table in doc.tables:
            rows_to_remove = []
            for row in table.rows:
                # Check if this is the old static skills description row
                text = " ".join(cell.text for cell in row.cells)
                if ("Schweißer (MIG/MAG" in text and 
                    "Sicherer Umgang mit Schweißgeräten" in text):
                    rows_to_remove.append(row._tr)
            for tr in rows_to_remove:
                table._tbl.remove(tr)

    def _clean_paragraphs(self, doc: Document):
        """Remove Jinja2 conditional blocks from paragraphs."""
        jinja_patterns = [
            r'\{%.*?%\}',       # {% if ... %}, {% endif %}, {%tr for ... %}, {%tr endfor %}
            r'\{\{.*?\}\}',     # {{ variable }}
        ]
        combined = '|'.join(jinja_patterns)

        for p in doc.paragraphs:
            full_text = p.text
            if re.search(combined, full_text):
                # Clear the paragraph
                for run in p.runs:
                    run.text = ""

    def _set_cell_text(self, cell, text: str):
        """Set cell text while preserving formatting."""
        # Clear existing runs
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""

        if cell.paragraphs:
            first_p = cell.paragraphs[0]
            first_p.clear()
            run = first_p.add_run(str(text))
            run.font.name = self._font_name
            run.font.size = self._font_size


# ------------------------------------------------------------------
# Extraction prompt builder
# ------------------------------------------------------------------

WELDER_MASTER_PROMPT = """Du bist ein professioneller deutscher Recruiting-Assistent für Schweißerkandidaten. 
Du erhältst einen Lebenslauf (CV) und musst daraus ein strukturiertes Profil erstellen.

EXTRAKTIONSREGELN:

1. BERUFSERFAHRUNG (employment_history):
   - Format: MM.JJJJ – MM.JJJJ (links), Details (rechts)
   - Sortierung: ab dem 19. Lebensjahr, lückenlos bis heute
   - Keine Datums-Lücken: Zwischenräume mit "Zeitpuffer/Projektpause" oder "Weiterbildung" füllen
   - Für jede Position: Arbeitgeber + Ort/Land, Position, Tätigkeit
   - Tätigkeit: Schweißverfahren (111, 135, 136, 138, 141), Material, typische Bauteile, Blech-/Wandstärken

2. BILDUNGSEINRICHTUNG (higher_education):
   - Schul-/Berufsausbildung aus dem CV übernehmen
   - Wenn fehlend: realistische Schule in der Nähe des Geburtsortes eintragen
   - Zeitraum 15.-18. Lebensjahr, Bezug zu Metallbearbeitung bevorzugt

3. WEITERBILDUNG (further_training):
   - Alle Schweißkurse und Weiterbildungen (111, 135, 136, 138, 141)
   - Mit Zeitraum und Art (Grund-/Aufbaukurs)

4. SPRACHKENNTNISSE (language_skills):
   - Tabellenform: Sprache + Niveau (C2, B1, A2, "kommunikativ")
   - "communicative" = B1-B2 interpretieren

5. TECHNISCHE FÄHIGKEITEN (technical_skills) - WICHTIG:
   Bewertungslogik:
   - Expert: 8-10+ Jahre regelmäßige Anwendung, anspruchsvolle Projekte
   - Advanced: 3-8 Jahre, regelmäßig angewendet, selbstständig
   - Intermediate: 1-3 Jahre, unter Anleitung oder einfache Projekte
   - Basic: nur Grundkenntnisse, kurze Einsätze

   Fähigkeiten zu bewerten:
   - Handwerkliches Geschick
   - Installation
   - Demontage
   - Blechbearbeitung
   - Montieren von Stahlkonstruktionen
   - Sägen / Schleifen / Schrauben
   - Bohren
   - Schweißarbeiten frei Hand
   - Schweißarbeiten mit Schweißroboter
   - MAG Schweißen
   - WIG Schweißen

   BEISPIEL: 15+ Jahre MAG → "Expert", WIG nur selten → "Intermediate"

6. DATENSCHUTZ: Private Kontaktdaten entfernen (Telefon, E-Mail, Adresse).
   Nur Platzhalter für Verleiher/Personaldienstleister verwenden.

ANTWORT FORMAT:
Antworte NUR mit gültigem JSON. Kein Markdown, keine Erklärungen.

{
  "candidate": {"full_name": "...", "nationality": "...", "date_of_birth": "YYYY-MM-DD", "position": "..."},
  "job_posting": {"title": "Schweißer (MIG/MAG)", "ekp": "", "svs": "", "start_date": ""},
  "employment_history": [
    {"start_date": "MM.JJJJ", "end_date": "MM.JJJJ", "employer": "...", "position": "...", "duties": ["..."]}
  ],
  "education": {
    "higher_education": [{"years": "YYYY - YYYY", "institution": "...", "field": "..."}],
    "further_training": [{"years": "YYYY", "institution": "...", "field": "..."}]
  },
  "technical_skills": [
    {"name": "MAG Schweißen", "level": "Expert"}
  ],
  "language_skills": [
    {"language": "Deutsch", "level": "B1"}
  ]
}
"""


def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF or DOCX file."""
    from pathlib import Path
    full_path = Path(file_path)
    suffix = full_path.suffix.lower()

    if suffix == ".pdf":
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(full_path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)

    elif suffix in (".docx", ".doc"):
        from docx import Document as DocxDocument
        doc = DocxDocument(str(full_path))
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text.strip())
        return "\n".join(text_parts)

    elif suffix in (".jpg", ".jpeg", ".png"):
        import base64
        # Return a marker that the caller should use OCR/Vision
        with open(str(full_path), "rb") as f:
            return f"__IMAGE_FILE__:{full_path}"
    else:
        raise ValueError(f"Unsupported format: {suffix}")


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract text from file bytes (for uploaded files)."""
    import tempfile
    import os
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        return extract_text_from_file(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
